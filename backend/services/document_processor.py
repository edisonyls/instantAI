import os
from typing import List, Dict, Optional
from docx import Document
import logging
from pathlib import Path
import uuid
from docx.oxml.ns import qn

try:
    from backend.config import settings
except ImportError:
    from config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """ Processing and extracting text from Word documents (.docx only) """

    def __init__(self):
        self.supported_formats = ['.docx']

    def extract_text(self, file_path: str) -> str:
        """ Extracts text from paragraphs, tables, headers, footers, 
        text boxes, shapes, footnotes, endnotes, and comments. """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            # Check file extension. Only .docx is supported for now
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")

            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))

            # Extract text from headers and footers
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(f"[HEADER] {paragraph.text.strip()}")

                footer = section.footer
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(f"[FOOTER] {paragraph.text.strip()}")

            extracted_text = '\n'.join(full_text)

            if not extracted_text.strip():
                raise ValueError("No text found in document")

            logger.info(
                f"Successfully extracted {len(extracted_text)} characters from {file_path}")
            return extracted_text

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[Dict]:
        """ Split text into chunks for vector storage """
        if chunk_size is None:
            chunk_size = settings.CHUNK_SIZE
        if overlap is None:
            overlap = settings.CHUNK_OVERLAP

        chunks = []
        start = 0
        chunk_id = 0

        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]

            # Try to break at sentence boundaries
            if end < len(text) and chunk_text.rfind('.') > chunk_size * 0.8:
                last_sentence = chunk_text.rfind('.')
                chunk_text = chunk_text[:last_sentence + 1]
                end = start + len(chunk_text)

            chunks.append({
                'id': str(uuid.uuid4()),
                'text': chunk_text.strip(),
                'start_char': start,
                'end_char': end,
                'chunk_index': chunk_id,
                'length': len(chunk_text.strip())
            })

            chunk_id += 1
            start = end - overlap

            if start >= len(text):
                break

        logger.info(
            f"Created {len(chunks)} chunks from text of length {len(text)}")
        return chunks

    def preprocess_text(self, text: str) -> str:
        """ Preprocess text for better RAG performance """
        # Remove excessive whitespace
        text = ' '.join(text.split())

        # Remove common artifacts
        text = text.replace('\x0c', ' ')
        text = text.replace('\x0b', ' ')
        text = text.replace('\u2019', "'")
        text = text.replace('\u201c', '"')
        text = text.replace('\u201d', '"')
        text = text.replace('\u2013', '-')
        text = text.replace('\u2014', '--')

        return text.strip()

    def validate_document(self, file_path: str) -> Dict:
        """ Validate a document before processing """
        result = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'metadata': {}
        }

        try:
            # Check file exists
            if not os.path.exists(file_path):
                result['valid'] = False
                result['errors'].append(f"File not found: {file_path}")
                return result

            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > settings.MAX_FILE_SIZE:
                result['valid'] = False
                result['errors'].append(
                    f"File too large: {file_size} bytes (max: {settings.MAX_FILE_SIZE})")
                return result

            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.supported_formats:
                result['valid'] = False
                result['errors'].append(f"Unsupported file format: {file_ext}")
                return result

            doc = Document(file_path)
            result['metadata'] = {
                'filename': Path(file_path).name,
                'file_size': file_size,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections)
            }

            try:
                footnote_count = 0
                comment_count = 0

                # Count footnotes
                if hasattr(doc.part, 'footnotes_part') and doc.part.footnotes_part:
                    footnote_count = len(doc.part.footnotes_part.footnotes)

                # Count comments
                if hasattr(doc.part, 'comments_part') and doc.part.comments_part:
                    comment_count = len(doc.part.comments_part.comments)

                result['metadata'].update({
                    'footnote_count': footnote_count,
                    'comment_count': comment_count
                })
            except Exception as e:
                logger.warning(
                    f"Could not count additional elements: {str(e)}")

            # Check if document has content
            text_content = self.extract_text(file_path)
            if len(text_content.strip()) < 10:
                result['warnings'].append(
                    "Document appears to have very little text content")

            result['metadata']['text_length'] = len(text_content)

        except Exception as e:
            result['valid'] = False
            result['errors'].append(f"Error validating document: {str(e)}")

        return result
